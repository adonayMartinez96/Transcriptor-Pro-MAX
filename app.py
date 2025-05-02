import os
import time
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF
import whisper
import subprocess
import datetime
from threading import Thread
from dotenv import load_dotenv
import requests
import json

class TranscriberPro:
    def __init__(self, root):
        self.root = root
        self.root.title("Whisper Transcriber Pro MAX")
        self.root.geometry("900x750")  # Aumentado para el nuevo botón
        self.transcription_done = False
        self.current_transcription = ""
        
         # Configuración de rutas FFmpeg
        self.ffmpeg_path, self.ffprobe_path = self._find_ffmpeg_paths()
        if not self.ffmpeg_path:
            messagebox.showerror("Error Crítico", 
                            "FFmpeg no está instalado correctamente.\n\n"
                            "Por favor ejecute 'install.bat' como administrador primero.")
            root.after(100, root.destroy)
        
        # Cargar configuración desde .env
        load_dotenv()
        self.deepseek_config = {
            'api_key': os.getenv('DEEPSEEK_API_KEY'),
            'model': os.getenv('DEEPSEEK_MODEL', 'deepseek-chat'),
            'base_url': os.getenv('DEEPSEEK_BASE_URL', 'https://openrouter.ai/api/v1'),
            'temperature': float(os.getenv('DEEPSEEK_TEMPERATURE', 0.7)),
            'max_tokens': int(os.getenv('DEEPSEEK_MAX_TOKENS', 2000)),
            'summary_prompt': os.getenv('SUMMARY_PROMPT', 'Resume el siguiente texto:\n\n')
        }
        
        # Variables
        self.file_path = tk.StringVar()
        self.model_var = tk.StringVar(value="base")
        self.progress_var = tk.DoubleVar()
        self.status_var = tk.StringVar(value="Listo")
        self.export_var = tk.IntVar(value=1)  # TXT por defecto
        
        # Configuración
        self.setup_styles()
        self.create_widgets()
    
    def setup_styles(self):
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('TFrame', background='#f5f5f5')
        style.configure('TLabel', background='#f5f5f5', font=('Arial', 10))
        style.configure('TButton', font=('Arial', 10), padding=6)
        style.configure('Header.TLabel', font=('Arial', 16, 'bold'), foreground='#333')
        style.configure('Progress.Horizontal.TProgressbar', thickness=25, troughcolor='#ddd', background='#4CAF50')
        style.configure('Summary.TButton', font=('Arial', 10, 'bold'), foreground='#1a5276')
    
    def create_widgets(self):
        # Frame principal
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Header
        header = ttk.Label(main_frame, text="WHISPER TRANSCRIBER PRO", style='Header.TLabel')
        header.pack(pady=(0, 20))
        
        # Sección de archivo
        file_frame = ttk.Frame(main_frame)
        file_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(file_frame, text="Archivo multimedia:").pack(side=tk.LEFT)
        
        self.file_entry = ttk.Entry(file_frame, textvariable=self.file_path, width=60)
        self.file_entry.pack(side=tk.LEFT, padx=10, expand=True, fill=tk.X)
        
        ttk.Button(file_frame, text="Examinar", command=self.browse_file).pack(side=tk.LEFT)
        
        # Sección de configuración
        config_frame = ttk.LabelFrame(main_frame, text="Configuración", padding=10)
        config_frame.pack(fill=tk.X, pady=10)
        
        # Modelos Whisper
        ttk.Label(config_frame, text="Modelo:").grid(row=0, column=0, sticky=tk.W, pady=5)
        
        models = [
            ("Tiny (rápido)", "tiny"),
            ("Base (recomendado)", "base"),
            ("Small", "small"),
            ("Medium", "medium"),
            ("Large (preciso)", "large")
        ]
        
        for i, (text, model) in enumerate(models):
            ttk.Radiobutton(config_frame, text=text, variable=self.model_var, 
                          value=model).grid(row=i+1, column=0, sticky=tk.W, padx=20)
        
        # Formatos de exportación
        ttk.Label(config_frame, text="Exportar a:").grid(row=0, column=1, sticky=tk.W, pady=5)
        
        formats = [
            ("Texto (.txt)", 1),
            ("Word (.docx)", 2),
            ("PDF (.pdf)", 3),
            ("Todos los formatos", 4)
        ]
        
        for i, (text, fmt) in enumerate(formats):
            ttk.Checkbutton(config_frame, text=text, variable=self.export_var, 
                           onvalue=fmt).grid(row=i+1, column=1, sticky=tk.W, padx=20)
        
        # Barra de progreso
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=15)
        
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var,
                                          maximum=100, style='Progress.Horizontal.TProgressbar')
        self.progress_bar.pack(fill=tk.X)
        
        self.status_label = ttk.Label(progress_frame, textvariable=self.status_var)
        self.status_label.pack(pady=5)
        
        # Botones de acción
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(pady=10)
        
        self.transcribe_btn = ttk.Button(button_frame, text="TRANSCRIBIR", command=self.start_transcription)
        self.transcribe_btn.pack(side=tk.LEFT, padx=5)
        
        self.summary_btn = ttk.Button(button_frame, text="GENERAR RESUMEN (DeepSeek)", 
                                    command=self.generate_summary, style='Summary.TButton', state=tk.DISABLED)
        self.summary_btn.pack(side=tk.LEFT, padx=5)
        
        # Resultados
        result_frame = ttk.LabelFrame(main_frame, text="Resultado", padding=10)
        result_frame.pack(fill=tk.BOTH, expand=True)
        
        self.result_text = tk.Text(result_frame, wrap=tk.WORD, font=('Arial', 10))
        self.result_text.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(self.result_text)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.result_text.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.result_text.yview)
    
    def browse_file(self):
        filetypes = [
            ('Videos', '*.mp4 *.avi *.mov *.mkv *.webm'),
            ('Audio', '*.mp3 *.wav *.ogg *.flac'),
            ('Todos los archivos', '*.*')
        ]
        
        filename = filedialog.askopenfilename(
            title='Seleccionar archivo multimedia',
            initialdir=os.path.expanduser('~'),
            filetypes=filetypes
        )
        
        if filename:
            self.file_path.set(filename)
            self.summary_btn.config(state=tk.DISABLED)
    
    def start_transcription(self):
        if not self.file_path.get():
            messagebox.showwarning("Advertencia", "Seleccione un archivo primero")
            return
            
        # Deshabilitar controles
        for child in self.root.winfo_children():
            if isinstance(child, ttk.Button):
                child.config(state=tk.DISABLED)
        
        # Ejecutar en hilo separado
        Thread(target=self.run_transcription, daemon=True).start()
    
    def run_transcription(self):
        try:
            self.status_var.set("Iniciando...")
            self.update_progress(5)
            
            # Cargar modelo
            self.status_var.set("Cargando modelo Whisper...")
            model = whisper.load_model(self.model_var.get())
            self.update_progress(15)
            
            # Transcribir con progreso aproximado
            self.status_var.set("Procesando archivo...")
            result = self.transcribe_with_progress(model, self.file_path.get())
            
            # Formatear texto con timestamps
            formatted_text = ""
            for segment in result['segments']:
                mins, secs = divmod(int(segment['start']), 60)
                formatted_text += f"{mins}:{secs:02d}\n{segment['text'].strip()}\n\n"
            
            self.current_transcription = formatted_text
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, formatted_text)
            self.update_progress(98)
            
            # Exportar resultados
            base_name = os.path.splitext(self.file_path.get())[0]
            export_format = self.export_var.get()
            
            exported_files = []
            
            # TXT
            if export_format in [1, 4]:
                txt_file = f"{base_name}_transcripcion.txt"
                with open(txt_file, 'w', encoding='utf-8') as f:
                    f.write(formatted_text)
                exported_files.append(txt_file)
            
            # DOCX
            if export_format in [2, 4]:
                docx_file = f"{base_name}_transcripcion.docx"
                self.export_to_docx(formatted_text, docx_file)
                exported_files.append(docx_file)
            
            # PDF
            if export_format in [3, 4]:
                pdf_file = f"{base_name}_transcripcion.pdf"
                self.export_to_pdf(formatted_text, pdf_file)
                exported_files.append(pdf_file)
            
            self.status_var.set(f"¡Completado! Exportado a: {', '.join(exported_files)}")
            self.update_progress(100)
            
            # Habilitar botón de resumen
            self.summary_btn.config(state=tk.NORMAL)
            messagebox.showinfo("Éxito", "Transcripción completada y exportada")
            
        except Exception as e:
            self.status_var.set("Error")
            messagebox.showerror("Error", f"Error durante la transcripción:\n{str(e)}")
        finally:
            # Reactivar controles
            self.transcribe_btn.config(state=tk.NORMAL)
            self.update_progress(0)
    
    def generate_summary(self):
            if not self.current_transcription:
                messagebox.showwarning("Advertencia", "No hay transcripción para resumir")
                return
                
            if not self.deepseek_config['api_key']:
                messagebox.showerror("Error", "API key de DeepSeek no configurada. Verifique el archivo .env")
                return
                
            # Deshabilitar botones durante la generación
            self.transcribe_btn.config(state=tk.DISABLED)
            self.summary_btn.config(state=tk.DISABLED)
            
            # Ejecutar en hilo separado
            Thread(target=self._generate_summary_thread, daemon=True).start()
            
    def transcribe_with_progress(self, model, filepath):
        """Transcribe con barra de progreso aproximada"""
        self.transcription_done = False
        estimated_duration = self.get_audio_duration(filepath) or 60  # Default 60s si no se puede obtener
        
        def update_progress():
            start_time = time.time()
            while not self.transcription_done and getattr(self, 'running', True):
                elapsed = time.time() - start_time
                progress = min(80, (elapsed / (estimated_duration * 0.3)) * 100)
                self.update_progress(15 + progress)
                time.sleep(0.3)
        
        progress_thread = threading.Thread(target=update_progress, daemon=True)
        progress_thread.start()
        
        try:
            result = model.transcribe(filepath, verbose=False, fp16=False)
            return result
        finally:
            self.transcription_done = True
            progress_thread.join()
            
    def _find_ffmpeg(self):
        """Busca FFmpeg en ubicaciones posibles con verificación de funcionamiento"""
        # 1. Primero verifica en la ruta del sistema
        try:
            subprocess.run(['ffmpeg', '-version'], check=True, capture_output=True)
            return 'ffmpeg'  # Usa el comando global
        except:
            pass
        
        # 2. Busca en rutas locales comunes
        local_paths = [
            os.path.join('ffmpeg', 'ffmpeg-master-latest-win64-gpl', 'bin', 'ffmpeg.exe'),
            os.path.join('ffmpeg', 'bin', 'ffmpeg.exe'),
            'ffmpeg.exe'
        ]
        
        for path in local_paths:
            full_path = os.path.abspath(path)
            if os.path.exists(full_path):
                try:
                    subprocess.run([full_path, '-version'], check=True, capture_output=True)
                    return full_path
                except:
                    continue
        
        return None  
    
    def _find_ffmpeg_paths(self):
        """Busca FFmpeg y FFprobe en ubicaciones posibles"""
        # 1. Verificar en el PATH del sistema
        try:
            subprocess.run(['ffmpeg', '-version'], check=True, capture_output=True)
            subprocess.run(['ffprobe', '-version'], check=True, capture_output=True)
            return ('ffmpeg', 'ffprobe')  # Usar comandos globales
        except:
            pass
        
        # 2. Buscar en rutas locales
        base_paths = [
            os.path.join('ffmpeg', 'ffmpeg-master-latest-win64-gpl', 'bin'),
            os.path.join('ffmpeg', 'bin'),
            ''
        ]
        
        for base in base_paths:
            ffmpeg = os.path.abspath(os.path.join(base, 'ffmpeg.exe'))
            ffprobe = os.path.abspath(os.path.join(base, 'ffprobe.exe'))
            
            if os.path.exists(ffmpeg) and os.path.exists(ffprobe):
                try:
                    subprocess.run([ffmpeg, '-version'], check=True, capture_output=True)
                    subprocess.run([ffprobe, '-version'], check=True, capture_output=True)
                    return (ffmpeg, ffprobe)
                except:
                    continue
        
        return (None, None)      
            
    def get_audio_duration(self, filepath):
        """Versión robusta para obtener duración del audio"""
        if not self.ffprobe_path:
            return None  # Fallback a duración por defecto
        
        try:
            cmd = [
                self.ffprobe_path,
                '-v', 'error',
                '-show_entries', 'format=duration',
                '-of', 'default=noprint_wrappers=1:nokey=1',
                filepath
            ]
            
            result = subprocess.run(cmd,
                                capture_output=True,
                                text=True,
                                check=True,
                                creationflags=subprocess.CREATE_NO_WINDOW)
            return float(result.stdout)
        except Exception as e:
            print(f"Error al obtener duración: {str(e)}")
            return None  # Fallback a duración por defecto

    
    def _generate_summary_thread(self):
            try:
                self.status_var.set("Generando resumen con DeepSeek...")
                self.update_progress(10)
                
                # Verificar nuevamente que tenemos la API key (por si acaso)
                if not self.deepseek_config.get('api_key'):
                    raise ValueError("API key no configurada. Verifique el archivo .env")
                
                headers = {
                    "Authorization": f"Bearer {self.deepseek_config['api_key']}",
                    "HTTP-Referer": "https://localhost",  # Necesario para OpenRouter
                    "X-Title": "Whisper Transcriber Pro",  # Identificación de la app
                    "Content-Type": "application/json"
                }
                
                # Limitar el texto a los tokens máximos permitidos (dejando espacio para la respuesta)
                max_content_length = self.deepseek_config['max_tokens'] * 4  # Estimación aproximada
                truncated_content = self.current_transcription[:max_content_length]
                
                prompt = self.deepseek_config['summary_prompt'] + truncated_content
                
                payload = {
                    "model": self.deepseek_config['model'],
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": self.deepseek_config['temperature'],
                    "max_tokens": self.deepseek_config['max_tokens']
                }
                
                # Añadir timeout para evitar que se quede colgado
                response = requests.post(
                    f"{self.deepseek_config['base_url']}/chat/completions",
                    headers=headers,
                    data=json.dumps(payload),
                    timeout=60  # 60 segundos de timeout
                )
                
                if response.status_code != 200:
                    error_msg = response.json().get('error', {}).get('message', response.text)
                    raise Exception(f"Error en la API ({response.status_code}): {error_msg}")
                
                result = response.json()
                summary = result['choices'][0]['message']['content']
                
                # Mostrar resumen en el hilo principal
                self.root.after(0, lambda: self.show_summary(summary))
                
                # Exportar resumen en el hilo principal
                self.root.after(0, self._export_summary, summary)
                
                self.root.after(0, lambda: self.status_var.set("Resumen generado con éxito"))
                self.update_progress(100)
                
            except requests.exceptions.Timeout:
                error_msg = "Tiempo de espera agotado. El servidor no respondió a tiempo."
                self.root.after(0, lambda: messagebox.showerror("Error", error_msg))
                self.status_var.set("Error: Timeout")
            except Exception as e:
                error_msg = str(e)
                self.root.after(0, lambda: messagebox.showerror("Error", f"Error al generar el resumen:\n{error_msg}"))
                self.status_var.set("Error al generar resumen")
            finally:
                self.root.after(0, self._enable_buttons)
                self.update_progress(0)

    def _export_summary(self, summary):
        """Maneja la exportación del resumen en el hilo principal"""
        try:
            base_name = os.path.splitext(self.file_path.get())[0]
            export_format = self.export_var.get()
            exported_files = []
            
            # TXT
            if export_format in [1, 4]:
                txt_file = f"{base_name}_resumen.txt"
                with open(txt_file, 'w', encoding='utf-8') as f:
                    f.write(summary)
                exported_files.append(txt_file)
            
            # DOCX
            if export_format in [2, 4]:
                docx_file = f"{base_name}_resumen.docx"
                self.export_to_docx(summary, docx_file)
                exported_files.append(docx_file)
            
            # PDF
            if export_format in [3, 4]:
                pdf_file = f"{base_name}_resumen.pdf"
                self.export_to_pdf(summary, pdf_file)
                exported_files.append(pdf_file)
            
            if exported_files:
                self.status_var.set(f"Resumen exportado a: {', '.join(exported_files)}")
        except Exception as e:
            messagebox.showerror("Error", f"Error al exportar el resumen:\n{str(e)}")
            
    def export_to_docx(self, text, filename):
        """Exporta el texto a un documento Word (.docx)"""
        try:
            doc = Document()
            
            # Configurar estilos del documento
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            
            # Agregar título
            doc.add_heading('Transcripción de Audio/Video', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Agregar contenido con formato
            for block in text.split('\n\n'):
                if block.strip():
                    lines = block.split('\n')
                    if len(lines) >= 2:
                        # Agregar timestamp como subtítulo
                        doc.add_paragraph(lines[0], style='Heading 2')
                        # Agregar texto de la transcripción
                        doc.add_paragraph('\n'.join(lines[1:]))
            
            doc.save(filename)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo exportar a DOCX:\n{str(e)}")
            raise

    def export_to_pdf(self, text, filename):
        """Exporta el texto a un documento PDF"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=11)
            
            # Título
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, txt="Transcripción de Audio/Video", ln=1, align='C')
            pdf.ln(10)
            
            # Contenido
            pdf.set_font("Arial", size=11)
            
            for block in text.split('\n\n'):
                if block.strip():
                    lines = block.split('\n')
                    if len(lines) >= 2:
                        # Timestamp en negrita
                        pdf.set_font("", 'B')
                        pdf.cell(0, 6, txt=lines[0], ln=1)
                        # Texto normal
                        pdf.set_font("", '')
                        pdf.multi_cell(0, 6, txt='\n'.join(lines[1:]))
                        pdf.ln(3)
            
            pdf.output(filename)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo exportar a PDF:\n{str(e)}")
            raise

    def _enable_buttons(self):
        """Reactivar botones en el hilo principal"""
        self.transcribe_btn.config(state=tk.NORMAL)
        self.summary_btn.config(state=tk.NORMAL)
    
    def show_summary(self, summary):
        top = tk.Toplevel(self.root)
        top.title("Resumen generado por DeepSeek")
        top.geometry("800x600")
        
        text_frame = ttk.Frame(top)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        text = tk.Text(text_frame, wrap=tk.WORD, font=('Arial', 11))
        text.pack(fill=tk.BOTH, expand=True)
        text.insert(tk.END, summary)
        
        scrollbar = ttk.Scrollbar(text)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        text.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=text.yview)
        
        btn_frame = ttk.Frame(top)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Cerrar", command=top.destroy).pack()
    
    def update_progress(self, value):
        self.progress_var.set(value)
        self.root.update_idletasks()

if __name__ == "__main__":
    root = tk.Tk()
    app = TranscriberPro(root)
    root.mainloop()
